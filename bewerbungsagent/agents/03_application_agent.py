import json
import os
import sys
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
from datetime import date

import requests
from bs4 import BeautifulSoup
from docx import Document
import anthropic

sys.path.append(os.path.join(os.path.dirname(__file__), ".."))

from utils.ollama_client import ask_ollama
from utils.excel_handler import add_bewerbung

# Pfade
PROFIL_PFAD = os.path.join(os.path.dirname(__file__), "..", "data", "profil.json")
ANSCHREIBEN_ORDNER = os.path.join(os.path.dirname(__file__), "..", "data", "anschreiben")

# E-Mail-Konfiguration (aus Umgebungsvariablen lesen)
SMTP_HOST = os.environ.get("SMTP_HOST", "smtp.gmail.com")
SMTP_PORT = int(os.environ.get("SMTP_PORT", "587"))
SMTP_USER = os.environ.get("SMTP_USER", "")
SMTP_PASSWORT = os.environ.get("SMTP_PASSWORT", "")


def find_jobs(jobtitel, ort, max_ergebnisse=10):
    """Sucht Jobs auf Indeed.de und gibt eine Liste von Job-Dicts zurück."""

    # Suchbegriffe URL-kodieren
    such_url = f"https://de.indeed.com/jobs?q={requests.utils.quote(jobtitel)}&l={requests.utils.quote(ort)}"

    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
    }

    try:
        response = requests.get(such_url, headers=headers, timeout=30)
        response.raise_for_status()
    except requests.exceptions.RequestException as e:
        print(f"Fehler beim Abrufen der Job-Suchergebnisse: {e}")
        return []

    # HTML parsen
    soup = BeautifulSoup(response.text, "html.parser")
    jobs = []

    # Indeed-Stellenanzeigen auslesen
    for karte in soup.find_all("div", class_="job_seen_beacon")[:max_ergebnisse]:
        try:
            # Titel
            titel_el = karte.find("h2", class_="jobTitle")
            titel = titel_el.get_text(strip=True) if titel_el else "Unbekannt"

            # Unternehmen
            firma_el = karte.find("span", {"data-testid": "company-name"})
            firma = firma_el.get_text(strip=True) if firma_el else "Unbekannt"

            # Ort
            ort_el = karte.find("div", {"data-testid": "text-location"})
            job_ort = ort_el.get_text(strip=True) if ort_el else ort

            # Kurzbeschreibung
            beschreibung_el = karte.find("div", class_="job-snippet")
            beschreibung = beschreibung_el.get_text(strip=True) if beschreibung_el else ""

            # URL
            link_el = karte.find("a", class_="jcs-JobTitle")
            job_url = "https://de.indeed.com" + link_el["href"] if link_el and link_el.get("href") else ""

            jobs.append({
                "titel": titel,
                "unternehmen": firma,
                "ort": job_ort,
                "beschreibung": beschreibung,
                "url": job_url
            })
        except Exception:
            # Einzelne fehlerhafte Karte überspringen
            continue

    print(f"{len(jobs)} Jobs gefunden für '{jobtitel}' in '{ort}'")
    return jobs


def _lade_profil():
    """Lädt das Benutzerprofil aus profil.json."""

    if not os.path.exists(PROFIL_PFAD):
        print("Kein Profil gefunden. Bitte zuerst Skills eingeben.")
        return None

    with open(PROFIL_PFAD, "r", encoding="utf-8") as f:
        return json.load(f)


def match_profile(job_dict):
    """Vergleicht ein Stellenangebot mit dem Benutzerprofil und gibt einen Match-Score zurück."""

    profil = _lade_profil()
    if profil is None:
        return None

    prompt = f"""Vergleiche diese Stellenanzeige mit dem Skill-Profil.
Antworte NUR mit JSON, kein anderer Text:
{{
  "match_score": 0,
  "vorhandene_skills": [],
  "fehlende_skills": [],
  "lernbare_skills": []
}}

Stellenanzeige:
{json.dumps(job_dict, ensure_ascii=False)}

Skill-Profil:
{json.dumps(profil["skills"], ensure_ascii=False)}"""

    antwort = ask_ollama(prompt)

    if antwort is None:
        return None

    # JSON extrahieren
    try:
        start = antwort.find("{")
        ende = antwort.rfind("}") + 1
        if start == -1 or ende == 0:
            print("Kein gültiges JSON vom LLM erhalten.")
            return None

        ergebnis = json.loads(antwort[start:ende])

        # Match-Score auf gültigen Bereich begrenzen
        ergebnis["match_score"] = max(0, min(100, int(ergebnis.get("match_score", 0))))
        return ergebnis

    except (json.JSONDecodeError, ValueError):
        print("Fehler beim Parsen der LLM-Antwort.")
        return None


def generate_cover_letter(job_dict, profil):
    """Erstellt ein Anschreiben mit der Claude API und speichert es als DOCX."""

    # Prompt für Claude
    prompt = f"""Schreibe ein professionelles deutsches Anschreiben für folgende Stelle.
Nutze konkrete Beispiele aus dem Skill-Profil.
Passe den Ton an den Unternehmenstyp an (Startup = locker/modern, Konzern = förmlich).

Stelle:
{json.dumps(job_dict, ensure_ascii=False, indent=2)}

Skill-Profil:
{json.dumps(profil, ensure_ascii=False, indent=2)}

Das Anschreiben soll haben: Betreff, Einleitung, Hauptteil mit konkreten Beispielen, Schluss mit Handlungsaufforderung."""

    # Claude API aufrufen
    client = anthropic.Anthropic()
    nachricht = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=1500,
        messages=[{"role": "user", "content": prompt}]
    )
    anschreiben_text = nachricht.content[0].text

    # Als DOCX speichern
    os.makedirs(ANSCHREIBEN_ORDNER, exist_ok=True)
    dateiname = f"Anschreiben_{job_dict['unternehmen'].replace(' ', '_')}_{date.today().isoformat()}.docx"
    dateipfad = os.path.join(ANSCHREIBEN_ORDNER, dateiname)

    doc = Document()
    doc.add_heading(f"Bewerbung: {job_dict['titel']}", level=1)

    # Text absatzweise einfügen
    for absatz in anschreiben_text.split("\n\n"):
        if absatz.strip():
            doc.add_paragraph(absatz.strip())

    doc.save(dateipfad)
    print(f"Anschreiben gespeichert: {dateipfad}")
    return dateipfad


def send_application(empfaenger_email, anschreiben_path, job_dict=None):
    """Sendet die Bewerbung per E-Mail — fragt zuerst nach Bestätigung."""

    # Immer zuerst Benutzer fragen
    print(f"\nBewerbung senden an: {empfaenger_email}")
    print(f"Anhang: {anschreiben_path}")
    bestaetigung = input("Wirklich senden? (ja/nein): ").strip().lower()

    if bestaetigung != "ja":
        print("Versand abgebrochen.")
        return False

    # E-Mail-Zugangsdaten prüfen
    if not SMTP_USER or not SMTP_PASSWORT:
        print("SMTP_USER und SMTP_PASSWORT Umgebungsvariablen sind nicht gesetzt.")
        return False

    # E-Mail zusammenbauen
    nachricht = MIMEMultipart()
    nachricht["From"] = SMTP_USER
    nachricht["To"] = empfaenger_email
    nachricht["Subject"] = f"Bewerbung: {job_dict['titel'] if job_dict else 'Stelle'}"

    # Begleittext
    nachricht.attach(MIMEText("Sehr geehrte Damen und Herren,\n\nim Anhang finden Sie meine Bewerbungsunterlagen.\n\nMit freundlichen Grüßen", "plain", "utf-8"))

    # Anhang hinzufügen
    if os.path.exists(anschreiben_path):
        with open(anschreiben_path, "rb") as f:
            teil = MIMEBase("application", "octet-stream")
            teil.set_payload(f.read())
            encoders.encode_base64(teil)
            teil.add_header("Content-Disposition", f"attachment; filename={os.path.basename(anschreiben_path)}")
            nachricht.attach(teil)

    # Versenden
    try:
        with smtplib.SMTP(SMTP_HOST, SMTP_PORT) as server:
            server.starttls()
            server.login(SMTP_USER, SMTP_PASSWORT)
            server.sendmail(SMTP_USER, empfaenger_email, nachricht.as_string())

        print("Bewerbung erfolgreich gesendet!")

        # In Excel-Protokoll eintragen
        if job_dict:
            add_bewerbung(
                titel=job_dict.get("titel", ""),
                unternehmen=job_dict.get("unternehmen", ""),
                match_score=job_dict.get("match_score", 0),
                status="Gesendet",
                url=job_dict.get("url", "")
            )
        return True

    except smtplib.SMTPException as e:
        print(f"Fehler beim E-Mail-Versand: {e}")
        return False


# Direkt ausführen zum Testen
if __name__ == "__main__":
    jobs = find_jobs("Python Entwickler", "Berlin")
    for job in jobs[:3]:
        print(f"- {job['titel']} bei {job['unternehmen']}")